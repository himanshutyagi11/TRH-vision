import os
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

docx_path = os.path.join('media', 'materials', 'Month_2_Learning_89WNej5.docx')
doc = Document(docx_path)

def inspect_element(p, location):
    if "Relation (Table)" in p.text:
        print(f"Match found in {location}:")
        print(f"  Text: {p.text}")
        txbx = p._p.findall('.//' + qn('w:txbxContent'))
        print(f"  Contains w:txbxContent: {len(txbx) > 0}")

# Search direct paragraphs
for i, p in enumerate(doc.paragraphs):
    inspect_element(p, f"doc.paragraphs[{i}]")

# Search tables
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for p_idx, p in enumerate(cell.paragraphs):
                inspect_element(p, f"doc.tables[{t_idx}].rows[{r_idx}].cells[{c_idx}].paragraphs[{p_idx}]")

